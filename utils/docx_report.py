"""Generate enterprise-styled VAPT reports as Microsoft Word (.docx) documents."""

from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from config.report_styles import (
    DEFAULT_SECTION_PREFIX,
    FILL_CRITICAL,
    FILL_HIGH,
    FILL_INFORMATIONAL,
    FILL_LOW,
    FILL_MEDIUM,
    FILL_METADATA_LABEL,
    FILL_NAVY_HEADER,
    FILL_WHITE,
    FONT_NAME,
    FONT_SIZE,
    FONT_SIZE_HEADING,
    PAGE_MARGIN,
    RGB_BLACK,
    RGB_WHITE,
)
from utils.schemas import AffectedEndpoint, ExecutiveSummary, ReportSection


def _severity_fill(severity: str) -> str:
    normalized = (severity or "").strip().upper()
    if "CRITICAL" in normalized:
        return FILL_CRITICAL
    if normalized == "HIGH" or normalized.startswith("HIGH"):
        return FILL_HIGH
    if normalized == "MEDIUM" or normalized.startswith("MEDIUM"):
        return FILL_MEDIUM
    if normalized == "LOW" or normalized.startswith("LOW"):
        return FILL_LOW
    return FILL_INFORMATIONAL


def _severity_text_color(fill_hex: str) -> RGBColor:
    if fill_hex in (FILL_CRITICAL, FILL_HIGH, FILL_NAVY_HEADER, FILL_METADATA_LABEL):
        return RGB_WHITE
    return RGB_BLACK


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "595959")
        borders.append(element)
    tc_pr.append(borders)


def _apply_run_font(run, *, bold: bool = False, color: Optional[RGBColor] = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _set_paragraph_font(paragraph, *, bold: bool = False, justify: bool = False) -> None:
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.space_before = Pt(0)
    if paragraph.runs:
        for run in paragraph.runs:
            _apply_run_font(run, bold=bold)
    else:
        run = paragraph.add_run()
        _apply_run_font(run, bold=bold)


def _write_styled_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    fill_hex: Optional[str] = None,
    font_color: Optional[RGBColor] = None,
    align_center: bool = False,
) -> None:
    if fill_hex:
        _set_cell_shading(cell, fill_hex)
    _set_cell_borders(cell)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    _apply_run_font(run, bold=bold, color=font_color or RGB_BLACK)


class VaptDocxReportBuilder:
    """Builds a client-deliverable .docx VAPT findings report."""

    def __init__(
        self,
        target: str,
        section_prefix: str = DEFAULT_SECTION_PREFIX,
    ) -> None:
        self.target = target
        self.section_prefix = section_prefix
        self.document = Document()
        self._configure_document()

    def _configure_document(self) -> None:
        section = self.document.sections[0]
        section.top_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN
        section.left_margin = PAGE_MARGIN
        section.right_margin = PAGE_MARGIN

        normal = self.document.styles["Normal"]
        normal.font.name = FONT_NAME
        normal.font.size = FONT_SIZE

    def add_cover_block(self, generated_at: str, findings_count: int) -> None:
        title = self.document.add_paragraph()
        run = title.add_run("Vulnerability Assessment & Penetration Testing Report")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(14)

        meta = self.document.add_paragraph()
        meta.add_run(f"Target: {self.target}\n").font.name = FONT_NAME
        meta.add_run(f"Generated: {generated_at}\n").font.name = FONT_NAME
        meta.add_run(f"Total findings: {findings_count}\n").font.name = FONT_NAME
        meta.paragraph_format.space_after = Pt(18)

    def add_executive_summary(self, summary: ExecutiveSummary) -> None:
        heading = self.document.add_paragraph()
        run = heading.add_run("Executive Summary")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_HEADING
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        risk = self.document.add_paragraph()
        risk.add_run(f"Overall risk level: {summary.overall_risk_level}").bold = True
        _set_paragraph_font(risk)

        body = self.document.add_paragraph(summary.executive_summary)
        _set_paragraph_font(body, justify=True)

        if summary.remediation_priorities:
            sub = self.document.add_paragraph()
            sub.add_run("Remediation priorities").bold = True
            _set_paragraph_font(sub)
            for item in summary.remediation_priorities:
                bullet = self.document.add_paragraph(item, style="List Bullet")
                _set_paragraph_font(bullet)

        self.document.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_finding(self, section: ReportSection, index: int) -> None:
        section_num = section.section_number or f"{self.section_prefix}.{index}"
        heading_text = f"{section_num} {section.title}"

        heading = self.document.add_paragraph()
        run = heading.add_run(heading_text)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_HEADING
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.space_after = Pt(8)

        self._add_metadata_table(section)
        self._add_description_section(section)
        if section.affected_endpoints:
            self._add_affected_endpoints_table(section.affected_endpoints)
        self._add_steps_to_reproduce(section.steps_to_reproduce)
        self._add_remediation(section.remediation_items())

    def _add_metadata_table(self, section: ReportSection) -> None:
        finding_id = section.finding_id or f"VAPT-{section.section_number or '001'}"
        rows_data = [
            ("ID", finding_id, False),
            ("Severity", section.severity.upper(), True),
            ("OWASP", section.owasp or "N/A", False),
            ("WSTG", section.wstg or "N/A", False),
        ]

        table = self.document.add_table(rows=len(rows_data), cols=2)
        table.autofit = False
        table.columns[0].width = Inches(1.4)
        table.columns[1].width = Inches(5.1)

        for row_idx, (label, value, is_severity) in enumerate(rows_data):
            label_cell = table.rows[row_idx].cells[0]
            value_cell = table.rows[row_idx].cells[1]

            _write_styled_cell(
                label_cell,
                label,
                bold=True,
                fill_hex=FILL_METADATA_LABEL,
                font_color=RGB_WHITE,
            )

            if is_severity:
                fill = _severity_fill(value)
                _write_styled_cell(
                    value_cell,
                    value,
                    bold=True,
                    fill_hex=fill,
                    font_color=_severity_text_color(fill),
                )
            else:
                _write_styled_cell(value_cell, value, fill_hex=FILL_WHITE)

        spacer = self.document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)

    def _add_description_section(self, section: ReportSection) -> None:
        sub = self.document.add_paragraph()
        sub.add_run("Description").bold = True
        _set_paragraph_font(sub)
        sub.paragraph_format.space_after = Pt(4)

        text = section.description or section.technical_description or ""
        body = self.document.add_paragraph(text)
        _set_paragraph_font(body, justify=True)
        body.paragraph_format.space_after = Pt(10)

    def _add_affected_endpoints_table(self, endpoints: List[AffectedEndpoint]) -> None:
        sub = self.document.add_paragraph()
        sub.add_run("Affected Endpoints").bold = True
        _set_paragraph_font(sub)
        sub.paragraph_format.space_after = Pt(4)

        headers = [
            "Endpoint",
            "Affected Low-Privileged Roles",
            "Impact",
            "Severity",
        ]
        table = self.document.add_table(rows=1 + len(endpoints), cols=4)
        table.autofit = False

        widths = [Inches(1.8), Inches(1.8), Inches(2.0), Inches(1.0)]
        for col_idx, width in enumerate(widths):
            table.columns[col_idx].width = width

        for col_idx, header in enumerate(headers):
            _write_styled_cell(
                table.rows[0].cells[col_idx],
                header,
                bold=True,
                fill_hex=FILL_NAVY_HEADER,
                font_color=RGB_WHITE,
            )

        for row_idx, endpoint in enumerate(endpoints, start=1):
            values = [
                endpoint.endpoint,
                endpoint.affected_roles,
                endpoint.impact,
                endpoint.severity,
            ]
            for col_idx, value in enumerate(values):
                cell = table.rows[row_idx].cells[col_idx]
                if col_idx == 3:
                    fill = _severity_fill(str(value))
                    _write_styled_cell(
                        cell,
                        str(value),
                        bold=True,
                        fill_hex=fill,
                        font_color=_severity_text_color(fill),
                        align_center=True,
                    )
                else:
                    _write_styled_cell(cell, str(value), fill_hex=FILL_WHITE)

        spacer = self.document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(10)

    def _add_steps_to_reproduce(self, steps: List[str]) -> None:
        sub = self.document.add_paragraph()
        sub.add_run("Steps to Reproduce").bold = True
        _set_paragraph_font(sub)
        sub.paragraph_format.space_after = Pt(4)

        for step in steps:
            para = self.document.add_paragraph(step, style="List Number")
            para.paragraph_format.left_indent = Inches(0.25)
            _set_paragraph_font(para)

        spacer = self.document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(10)

    def _add_remediation(self, items: List[str]) -> None:
        sub = self.document.add_paragraph()
        sub.add_run("Remediation").bold = True
        _set_paragraph_font(sub)
        sub.paragraph_format.space_after = Pt(4)

        for item in items:
            para = self.document.add_paragraph(item, style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.25)
            _set_paragraph_font(para)

        spacer = self.document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(14)


def build_vapt_docx(
    target: str,
    sections: List[ReportSection],
    *,
    generated_at: str,
    executive_summary: Optional[ExecutiveSummary] = None,
    section_prefix: str = DEFAULT_SECTION_PREFIX,
) -> Document:
    builder = VaptDocxReportBuilder(target=target, section_prefix=section_prefix)
    builder.add_cover_block(generated_at, len(sections))

    if executive_summary:
        builder.add_executive_summary(executive_summary)

    for index, section in enumerate(sections, start=1):
        builder.add_finding(section, index)

    return builder.document
